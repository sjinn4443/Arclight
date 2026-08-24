from __future__ import annotations

import hashlib
import os
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from collections import Counter
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path
from zoneinfo import ZoneInfo

import pymupdf
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from .database import IpRow, ReportSnapshot, UsageWindow
from .geo import GeoLocation, GeoResolver
from .safe import safe_text

NAVY = "18324A"
TEAL = "2D6F73"
PALE = "EAF2F3"
GREY = "5C6873"
WHITE = "FFFFFF"
RED = "A43D3D"


@dataclass(frozen=True, slots=True)
class ReportArtifact:
    report_id: str
    display_name: str
    path: Path
    size: int
    sha256: str
    page_count: int
    validation_passed: bool
    validation_exceptions: tuple[str, ...]


def _set_cell_shading(cell, colour: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), colour)


def _set_cell_width(cell, inches: float) -> None:
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(inches * 1440)))
    width.set(qn("w:type"), "dxa")
    cell._tc.get_or_add_tcPr().append(width)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _page_number(paragraph) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def _format_timestamp(value: datetime | None, timezone: ZoneInfo) -> str:
    if value is None:
        return "Unavailable"
    if value.tzinfo is None:
        value = value.replace(tzinfo=UTC)
    return value.astimezone(timezone).strftime("%d %b %Y, %H:%M %Z")


def _configure_document(document: Document, timezone: ZoneInfo) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, colour in (
        ("Title", 26, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 12, TEAL),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(colour)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "ARCLIGHT  /  PRIVATE POSTGRESQL REPORT"
    header.style = styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    header.runs[0].font.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Private developer report  •  ")
    _page_number(footer)

    props = document.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    props.category = ""
    props.identifier = ""
    props.keywords = ""
    props.subject = ""
    props.title = ""
    props.language = "en-GB"
    props.created = datetime.now(timezone).replace(tzinfo=None)
    props.modified = datetime.now(timezone).replace(tzinfo=None)


def _add_status_banner(document: Document, passed: bool) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, PALE if passed else "F8E8E8")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "LINK VALIDATION PASSED — PROFILE AND LATEST-IP ROWS INCLUDED"
        if passed
        else "LINK VALIDATION FAILED — LINKED PROFILE/IP ROWS OMITTED"
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEAL if passed else RED)


def _add_usage_table(document: Document, usage: dict[str, UsageWindow]) -> None:
    headers = (
        "Period",
        "All profiles",
        "New",
        "Active",
        "Refreshes*",
        "IP events",
        "Unique IPs",
        "Countries",
    )
    labels = {"24_hours": "24 hours", "7_days": "7 days", "30_days": "30 days"}
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.05, 0.9, 0.65, 0.65, 0.85, 0.75, 0.75, 0.75)
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.rows[0].cells[index]
        _set_cell_width(cell, width)
        _set_cell_shading(cell, NAVY)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(7.5)
    _repeat_header(table.rows[0])

    for key in ("24_hours", "7_days", "30_days"):
        value = usage[key]
        row = table.add_row()
        values = (
            labels[key],
            value.total_profiles,
            value.new_profiles,
            value.active_profiles,
            value.refresh_count_on_active_profiles,
            value.ip_events,
            value.unique_ips,
            value.countries,
        )
        for cell, item, width in zip(row.cells, values, widths, strict=True):
            _set_cell_width(cell, width)
            cell.text = str(item)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
        _prevent_row_split(row)
    note = document.add_paragraph(
        "* Refresh count is cumulative for profiles active in the period because "
        "the source table does not store individual refresh-event timestamps."
    )
    note.style = document.styles["Caption"]


def _landscape_section(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    return section


def _add_table(
    document: Document,
    headers: tuple[str, ...],
    widths: tuple[float, ...],
    rows: list[tuple[str, ...]],
    *,
    font_size: float = 7.5,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.rows[0].cells[index]
        _set_cell_width(cell, width)
        _set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    _repeat_header(table.rows[0])

    for row_index, values in enumerate(rows):
        row = table.add_row()
        if row_index % 2:
            for cell in row.cells:
                _set_cell_shading(cell, "F5F8F9")
        for cell, value, width in zip(row.cells, values, widths, strict=True):
            _set_cell_width(cell, width)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
        _prevent_row_split(row)


def _enrich_history(
    history: tuple[IpRow, ...], resolver: GeoResolver
) -> list[tuple[IpRow, GeoLocation]]:
    return [(row, resolver.resolve(row.ip, row.country_name)) for row in history]


def build_report_document(
    snapshot: ReportSnapshot,
    *,
    timezone_name: str,
    geoip_db_path: Path | None,
) -> Document:
    timezone = ZoneInfo(timezone_name)
    document = Document()
    _configure_document(document, timezone)

    eyebrow = document.add_paragraph("WEEKLY PRIVATE DATA REPORT")
    eyebrow.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    eyebrow.runs[0].font.bold = True
    eyebrow.paragraph_format.space_after = Pt(3)
    document.add_heading("Arclight PostgreSQL report", 0)
    subtitle = document.add_paragraph(
        f"Generated {_format_timestamp(snapshot.as_of, timezone)}  •  Europe/London"
    )
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(GREY)
    _add_status_banner(document, snapshot.validation.passed)

    document.add_heading("Usage at a glance", level=1)
    _add_usage_table(document, snapshot.usage)

    document.add_heading("Data validation", level=1)
    validation_rows = [
        ("Profile rows", str(snapshot.user_count)),
        ("Latest-IP rows", str(snapshot.latest_ip_count)),
        ("IP history rows", str(snapshot.ip_history_count)),
        (
            "Pairing decision",
            "Included" if snapshot.validation.passed else "Omitted",
        ),
        (
            "Exceptions",
            ", ".join(snapshot.validation.exceptions) or "None",
        ),
        (
            "GeoIP method",
            "Offline city database plus stored country"
            if geoip_db_path and geoip_db_path.is_file()
            else "Stored country only; city marked unavailable",
        ),
    ]
    _add_table(
        document,
        ("Check", "Result"),
        (1.7, 5.4),
        validation_rows,
        font_size=8.5,
    )
    paragraph = document.add_paragraph(
        "Pairing follows row N to row N only after equal counts, unique profile IDs, "
        "expected view shapes and repeated deterministic ordering all pass."
    )
    paragraph.style = document.styles["Caption"]

    with GeoResolver(geoip_db_path) as resolver:
        history = _enrich_history(snapshot.ip_history, resolver)

        country_counts = Counter(location.country for _, location in history)
        city_counts = Counter(location.city for _, location in history)
        document.add_heading("Location summary", level=1)
        summary_rows: list[tuple[str, str, str]] = []
        for label, counts in (("Country", country_counts), ("City", city_counts)):
            for value, count in counts.most_common(20):
                summary_rows.append((label, value, str(count)))
        if not summary_rows:
            summary_rows.append(("Location", "No IP history rows", "0"))
        _add_table(
            document,
            ("Level", "Location", "Events"),
            (1.0, 4.9, 0.9),
            summary_rows,
            font_size=8.5,
        )

        _landscape_section(document)
        document.add_heading("Latest profile and IP rows", level=1)
        if snapshot.validation.passed:
            linked_rows: list[tuple[str, ...]] = []
            for user, ip_row in zip(snapshot.users, snapshot.latest_ips, strict=True):
                location = resolver.resolve(ip_row.ip, ip_row.country_name)
                linked_rows.append(
                    (
                        safe_text(user.name, limit=120) or "—",
                        safe_text(user.contact, limit=180) or "—",
                        safe_text(user.profile_id, limit=180),
                        safe_text(ip_row.ip, limit=64),
                        location.city,
                        location.country,
                        _format_timestamp(user.first_seen, timezone),
                        _format_timestamp(user.last_seen, timezone),
                        str(user.refresh_count),
                    )
                )
            _add_table(
                document,
                (
                    "Name",
                    "Contact",
                    "Profile ID",
                    "Raw IP",
                    "City",
                    "Country",
                    "First seen",
                    "Last seen",
                    "Refresh",
                ),
                (1.05, 1.35, 2.25, 1.15, 0.8, 0.85, 1.15, 1.15, 0.55),
                linked_rows,
                font_size=6.8,
            )
        else:
            document.add_paragraph(
                "No linked name, contact, profile ID and raw-IP rows are included "
                "because the row-pairing validation failed."
            )

        _landscape_section(document)
        document.add_heading("IP history", level=1)
        document.add_paragraph(
            "These rows are chronological and intentionally unlinked to profiles "
            "because the source IP table has no profile identifier."
        )
        history_rows = [
            (
                safe_text(row.ip, limit=64),
                location.city,
                location.country,
                _format_timestamp(row.ts, timezone),
            )
            for row, location in history
        ]
        if not history_rows:
            history_rows = [("—", "—", "—", "No IP history rows")]
        _add_table(
            document,
            ("Raw IP", "City", "Country", "Observed"),
            (2.0, 2.4, 2.4, 2.4),
            history_rows,
            font_size=7.3,
        )

    return document


def _scrub_docx_metadata(path: Path) -> None:
    core_ns = {
        "dc": "http://purl.org/dc/elements/1.1/",
        "dcterms": "http://purl.org/dc/terms/",
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    }
    app_ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}
    replacement = path.with_suffix(".scrubbed.docx")
    with (
        zipfile.ZipFile(path, "r") as source,
        zipfile.ZipFile(replacement, "w", compression=zipfile.ZIP_DEFLATED) as target,
    ):
        for item in source.infolist():
            if item.filename == "docProps/custom.xml":
                continue
            data = source.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = etree.fromstring(data)
                for xpath in (
                    "//dc:creator",
                    "//dc:subject",
                    "//dc:title",
                    "//cp:lastModifiedBy",
                    "//cp:keywords",
                    "//cp:category",
                    "//cp:contentStatus",
                    "//cp:lastPrinted",
                    "//dcterms:created",
                    "//dcterms:modified",
                ):
                    for node in root.xpath(xpath, namespaces=core_ns):
                        node.getparent().remove(node)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "docProps/app.xml":
                root = etree.fromstring(data)
                for xpath in ("//ep:Company", "//ep:Manager"):
                    for node in root.xpath(xpath, namespaces=app_ns):
                        node.text = ""
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            target.writestr(item, data)
    os.chmod(replacement, 0o600)
    os.replace(replacement, path)


def _render_and_inspect(docx_path: Path, soffice_binary: str) -> int:
    staging = Path(tempfile.mkdtemp(prefix="arclight-render-"))
    profile = staging / "lo-profile"
    profile.mkdir()
    try:
        completed = subprocess.run(
            [
                soffice_binary,
                f"-env:UserInstallation={profile.as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(staging),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=240,
        )
        pdf_path = staging / f"{docx_path.stem}.pdf"
        if completed.returncode != 0 or not pdf_path.is_file():
            raise RuntimeError("docx_render_failed")

        with pymupdf.open(pdf_path) as pdf:
            if pdf.page_count < 1:
                raise RuntimeError("docx_render_empty")
            for page_index, page in enumerate(pdf):
                text = page.get_text("text").strip()
                if len(text) < 8:
                    raise RuntimeError("docx_page_blank")
                tolerance = 1.5
                for block in page.get_text("blocks"):
                    x0, y0, x1, y1 = block[:4]
                    if (
                        x0 < -tolerance
                        or y0 < -tolerance
                        or x1 > page.rect.width + tolerance
                        or y1 > page.rect.height + tolerance
                    ):
                        raise RuntimeError("docx_content_clipped")
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False)
                page_png = staging / f"page-{page_index + 1:04d}.png"
                pixmap.save(page_png)
                if not page_png.is_file() or page_png.stat().st_size < 1_000:
                    raise RuntimeError("docx_page_render_failed")
            return pdf.page_count
    finally:
        shutil.rmtree(staging, ignore_errors=True)


def _cleanup_expired_reports(output_dir: Path, retention_hours: int) -> None:
    cutoff = datetime.now(UTC) - timedelta(hours=retention_hours)
    for path in output_dir.glob("Arclight_Postgres_Report_*.docx"):
        try:
            modified = datetime.fromtimestamp(path.stat().st_mtime, UTC)
            if modified < cutoff:
                path.unlink()
        except OSError:
            continue


def _finalise_document(
    document: Document,
    *,
    output_dir: Path,
    retention_hours: int,
    report_date: str,
    soffice_binary: str,
    validation_passed: bool,
    validation_exceptions: tuple[str, ...],
) -> ReportArtifact:
    output_dir.mkdir(mode=0o700, parents=True, exist_ok=True)
    os.chmod(output_dir, 0o700)
    report_id = str(uuid.uuid4())
    display_name = f"Arclight_Postgres_Report_{report_date}.docx"
    final_path = output_dir / f"Arclight_Postgres_Report_{report_date}_{report_id}.docx"
    staging = Path(tempfile.mkdtemp(prefix="arclight-docx-"))
    try:
        staged_path = staging / display_name
        document.save(staged_path)
        os.chmod(staged_path, 0o600)
        _scrub_docx_metadata(staged_path)
        page_count = _render_and_inspect(staged_path, soffice_binary)
        os.replace(staged_path, final_path)
        os.chmod(final_path, 0o600)
    finally:
        shutil.rmtree(staging, ignore_errors=True)
    _cleanup_expired_reports(output_dir, retention_hours)
    digest = hashlib.sha256(final_path.read_bytes()).hexdigest()
    return ReportArtifact(
        report_id=report_id,
        display_name=display_name,
        path=final_path,
        size=final_path.stat().st_size,
        sha256=digest,
        page_count=page_count,
        validation_passed=validation_passed,
        validation_exceptions=validation_exceptions,
    )


def generate_report(
    snapshot: ReportSnapshot,
    *,
    output_dir: Path,
    retention_hours: int,
    timezone_name: str,
    geoip_db_path: Path | None,
    soffice_binary: str,
) -> ReportArtifact:
    timezone = ZoneInfo(timezone_name)
    report_date = snapshot.as_of.astimezone(timezone).strftime("%Y-%m-%d")
    _cleanup_expired_reports(output_dir, retention_hours)
    document = build_report_document(
        snapshot,
        timezone_name=timezone_name,
        geoip_db_path=geoip_db_path,
    )
    return _finalise_document(
        document,
        output_dir=output_dir,
        retention_hours=retention_hours,
        report_date=report_date,
        soffice_binary=soffice_binary,
        validation_passed=snapshot.validation.passed,
        validation_exceptions=snapshot.validation.exceptions,
    )


def generate_error_report(
    *,
    output_dir: Path,
    retention_hours: int,
    timezone_name: str,
    soffice_binary: str,
    error_code: str,
) -> ReportArtifact:
    timezone = ZoneInfo(timezone_name)
    now = datetime.now(timezone)
    document = Document()
    _configure_document(document, timezone)
    document.add_heading("Arclight PostgreSQL report", 0)
    document.add_paragraph(now.strftime("Generated %d %b %Y, %H:%M %Z"))
    document.add_heading("Data unavailable", level=1)
    document.add_paragraph(
        "The production PostgreSQL read failed. No personal data was retrieved or "
        "included in this report."
    )
    document.add_paragraph(f"Error code: {safe_text(error_code, limit=80)}")
    _cleanup_expired_reports(output_dir, retention_hours)
    return _finalise_document(
        document,
        output_dir=output_dir,
        retention_hours=retention_hours,
        report_date=now.strftime("%Y-%m-%d"),
        soffice_binary=soffice_binary,
        validation_passed=False,
        validation_exceptions=(error_code,),
    )


def find_report(output_dir: Path, report_id: str) -> Path | None:
    try:
        canonical = str(uuid.UUID(report_id))
    except ValueError:
        return None
    matches = list(output_dir.glob(f"Arclight_Postgres_Report_*_{canonical}.docx"))
    if len(matches) != 1:
        return None
    path = matches[0].resolve()
    try:
        path.relative_to(output_dir.resolve())
    except ValueError:
        return None
    return path if path.is_file() else None
